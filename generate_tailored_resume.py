"""
Rebuild tailored resume following the EXACT format/structure of the original resume.
Rules:
1. Follow exact template format and section positions from original resume
2. Keep all content - only modify/add where relevant to job description
3. Keep research projects (remove only if completely irrelevant to role)
4. Keep ValueLabs section as-is or enhanced (never remove points)
5. Update skills/summary/bullets strategically to match JD keywords
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
import os
import re


def create_tailored_resume(job_title, company, location, jd_keywords):
    """Create a tailored resume following exact original template structure."""
    
    doc = Document()
    
    # Set page margins (narrow like original)
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9.5)
    
    # ═══════════════════════════════════════════════════════════════════
    # HEADER - Name + Title + Contact (exactly as original)
    # ═══════════════════════════════════════════════════════════════════
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Mohammad Sheriff Mehmood')
    run.font.size = Pt(18)
    run.font.bold = True
    p.space_after = Pt(2)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Senior Data Scientist · NLP & Generative AI Engineer')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)
    p.space_after = Pt(4)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('+91 7999937693 | mdsheriff2702@gmail.com | LinkedIn | github.com/sheriff786 | Portfolio')
    run.font.size = Pt(9)
    p.space_after = Pt(8)
    
    # ═══════════════════════════════════════════════════════════════════
    # PROFESSIONAL SUMMARY (same position as original)
    # ═══════════════════════════════════════════════════════════════════
    
    add_section_heading(doc, 'PROFESSIONAL SUMMARY')
    
    # Tailored summary - keep core but add JD-relevant keywords
    summary = (
        'Senior Data Scientist & AI/ML Engineer with 6+ Years of experience. '
        'I build AI systems that move business metrics — not just demos. '
        'Specializing in NLP, LLMs, and production MLOps, I have shipped enterprise-grade GenAI systems '
        'that delivered $1M+ revenue uplift, cut churn by 40%, and accelerated model iteration by 60% '
        'across healthcare, finance, and retail. '
    )
    
    # Add JD-specific additions to summary
    if 'databricks' in company.lower():
        summary += (
            'Experienced in building production-grade GenAI applications including RAG, multi-agent systems, '
            'and fine-tuning with tools like LangChain, HuggingFace, and PyTorch. '
            'Proven track record as a technical advisor deploying AI solutions on AWS and GCP. '
        )
    elif 'airbnb' in company.lower():
        summary += (
            'Experienced in building scalable AI-powered applications with RAG systems and '
            'cross-functional collaboration across product, engineering, and operations teams. '
        )
    
    summary += 'Open-source author: published authentica on PyPI, actively used in production at Businessolver (Healthcare Benefits SaaS).'
    
    p = doc.add_paragraph(summary)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(9.5)
    
    # ═══════════════════════════════════════════════════════════════════
    # TECHNICAL SKILLS (same layout as original)
    # ═══════════════════════════════════════════════════════════════════
    
    add_section_heading(doc, 'TECHNICAL SKILLS')
    
    # Base skills (from original resume)
    skills_data = [
        ('Languages', 'Python (primary) — proficient in TensorFlow, PyTorch, Scikit-learn, NumPy, Pandas, OpenCV, YOLO, CNN'),
        ('GenAI & LLMs', 'LangChain · LangSmith · OpenAI API · Gemini · RAG · LLMOps · Fine-tuning · Prompt Engineering · Agentic AI · MCP · Transformer Models · Sentence-Transformers · CrewAI'),
        ('MLOps & Cloud', 'MLflow · Docker · Kubernetes · Airflow · GCP (Vertex AI) · AWS (SageMaker, ECS, CodePipeline) · CI/CD'),
        ('Databases', 'MongoDB · SQL · PostgreSQL · ChromaDB · FAISS · Pinecone'),
        ('Tools & Other', 'Tableau · Weights & Biases · FastAPI · Flask · Django · A/B Testing · REST API Design · Git'),
    ]
    
    # Add JD-specific skills if not already present
    if 'databricks' in company.lower():
        skills_data[1] = ('GenAI & LLMs', 'LangChain · LangSmith · OpenAI API · Gemini · RAG · LLMOps · Fine-tuning · Prompt Engineering · Agentic AI · MCP · Transformer Models · Sentence-Transformers · CrewAI · HuggingFace · DSPy · Multi-agent Systems · Text2SQL')
        skills_data[2] = ('MLOps & Cloud', 'MLflow · Docker · Kubernetes · Airflow · GCP (Vertex AI) · AWS (SageMaker, ECS, CodePipeline) · Azure · CI/CD · Apache Spark')
    
    for category, items in skills_data:
        p = doc.add_paragraph()
        run = p.add_run(f'{category}: ')
        run.font.bold = True
        run.font.size = Pt(9.5)
        run = p.add_run(items)
        run.font.size = Pt(9.5)
        p.paragraph_format.space_after = Pt(2)
    
    # ═══════════════════════════════════════════════════════════════════
    # PROFESSIONAL EXPERIENCE (exact same structure as original)
    # ═══════════════════════════════════════════════════════════════════
    
    add_section_heading(doc, 'PROFESSIONAL EXPERIENCE')
    
    # --- ValueLabs (Mar 2025 - Present) ---
    add_company_header(doc, 'ValueLabs · Hyderabad, India', 'Senior Data Scientist | AI/ML Engineer', 'Mar 2025 – Present')
    
    # Keep and enhance ValueLabs bullets (user said don't remove!)
    valuelabs_intro = (
        'Embedded with Generative AI practice, building production AI systems for enterprise clients.'
    )
    p = doc.add_paragraph(valuelabs_intro)
    p.paragraph_format.space_before = Pt(2)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.italic = True
    
    valuelabs_bullets = [
        'Developing and deploying production-grade GenAI applications leveraging RAG, multi-agent systems, and LLM fine-tuning for enterprise knowledge management solutions.',
        'Building end-to-end LLMOps pipelines using LangChain, OpenAI, and vector databases (FAISS, ChromaDB) for scalable AI-powered applications.',
        'Serving as technical advisor to cross-functional teams, translating complex AI concepts into actionable business solutions driving measurable ROI.',
    ]
    
    for bullet in valuelabs_bullets:
        add_bullet(doc, bullet)
    
    # --- LTIMindtree (Jan 2024 - Mar 2025) ---
    add_company_header(doc, 'LTIMindtree · Pune, India', 'Senior Data Scientist – NLP Engineer', 'Jan 2024 – Mar 2025')
    
    ltim_intro = (
        'Embedded with Businessolver\'s Data Science team as a full-stack ML Engineer; designed and shipped an automated '
        'document analysis pipeline leveraging computer vision, VLMs, and AWS-native tooling to process healthcare documents at scale.'
    )
    p = doc.add_paragraph(ltim_intro)
    p.paragraph_format.space_before = Pt(2)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.italic = True
    
    ltim_bullets = [
        'GenAI Chatbot & RAG System: Architected production-grade RAG chatbot (Gemini + FAISS Vector DB) serving 15+ team members with 90% query accuracy and sub-2s response time for enterprise knowledge retrieval.',
        'LLMOps Pipeline: Built end-to-end LLMOps pipeline using OpenAI, LangChain, ChromaDB, and sentence-transformers cutting model iteration time by 60% and reducing deployment overhead by 50%.',
        'Healthcare Revenue Model: Deployed anomaly-detection targeting model generating $1M+ in annual healthcare revenue at 90% precision; reduced false positives by 20% via class-weighted ensemble (Logistic Regression + Random Forest).',
        'Workflow Orchestration: Designed Airflow DAGs to automate multi-step ML pipelines, cutting end-to-end processing time by 40% and improving delivery velocity across a 7-member team.',
        'Semantic NLP (BioBERT): Generated vector embeddings for ICD-9 and medical procedure codes using BioBERT, improving semantic search accuracy by 18% over TF-IDF baseline.',
        'Document Analysis Pipeline: Designed an end-to-end multi-modal pipeline on AWS (Comprehend, Textract, S3, ECS) to classify and extract data from 47,000+ healthcare documents, significantly cutting manual claim validation time.',
        'Query-Based Document Understanding: Built a custom QA system integrating visual features from Textract OCR with textual embeddings across varied layouts; deployed inference on AWS Lambda achieving 93% extraction precision.',
        'Annotation & Monitoring UI: Built an interactive React UI with REST API for data annotation/validation, real-time model metric visualisation, performance monitoring, and data-driven threshold tuning.',
        'Stacking Ensemble Model: Trained and tuned a stacking ensemble (Random Forest, XGBoost, LightGBM) via cross-validation achieving 0.78 F1 and 0.81 AUC; applied SHAP and Gini impurity for model explainability.',
    ]
    
    for bullet in ltim_bullets:
        add_bullet(doc, bullet)
    
    # --- Cognizant (Dec 2019 - Jan 2024) ---
    add_company_header(doc, 'Cognizant Technology Solutions · Chennai, India', 'Data Scientist', 'Dec 2019 – Jan 2024')
    
    cog_intro = 'Joined as Programmer Analyst → core ML systems, promoted to Data Scientist (leading cross-functional AI/NLP initiatives).'
    p = doc.add_paragraph(cog_intro)
    p.paragraph_format.space_before = Pt(2)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.italic = True
    
    cognizant_bullets = [
        'ML Model Development: Built and deployed supervised ML models improving prediction accuracy by 20% and saving $200K in operational costs; improved baselines by 5–10% using Neural Networks, ensemble methods and GridSearchCV hyperparameter tuning.',
        'AI Customer Support Platform: Led 7-member cross-functional team to build NLP-powered support application; reduced customer churn by 40% and generated $300K additional quarterly revenue through context-aware intelligent responses.',
        'Fraud Detection System: Reduced false-positive fraud alerts by 20% using Logistic Regression and Random Forest with class-weighted resampling — received company Innovation Award.',
        'Data Platform & BI: Unified 10+ heterogeneous data sources (databases, APIs, third-party feeds) into ML-ready pipelines; built Tableau dashboards tracking 20+ KPIs, improving cross-team reporting efficiency by 30%.',
    ]
    
    for bullet in cognizant_bullets:
        add_bullet(doc, bullet)
    
    # ═══════════════════════════════════════════════════════════════════
    # RESEARCH PROJECTS (keep all - remove only if irrelevant)
    # ═══════════════════════════════════════════════════════════════════
    
    add_section_heading(doc, 'RESEARCH PROJECTS')
    
    # All projects from original resume
    projects = [
        {
            'title': 'Abstractive Text Summarisation',
            'tech': 'Python · BERT · NLP · Transformers',
            'bullets': [
                'Applied BERT and NLP toolkits to summarise live chat context, improving agent efficiency by 35% and reducing customer response times by 25%.',
                'Gave agents instant access to previous chat history — reducing handle time and improving resolution quality.',
            ],
            'relevant_to': ['data scientist', 'nlp', 'ai engineer', 'ml engineer'],
        },
        {
            'title': 'End-to-End Topic Modelling & MLOps on AWS',
            'tech': 'Python · LDA · Docker · AWS ECS · MLflow · CodePipeline · EC2',
            'bullets': [
                'Built LDA topic modelling pipeline on AWS ECS with blue/green deployment — cut processing time by 40% and deployment time by 60%.',
                'Automated full CI/CD (CodeCommit, CodeBuild, CodeDeploy, CodePipeline); reduced manual effort by 70%. Used MLflow for model versioning; ensured high availability via EC2 + ECS load balancing.',
            ],
            'relevant_to': ['data scientist', 'ml engineer', 'mlops', 'ai engineer'],
        },
        {
            'title': 'Visa Approval Forecasting (US Immigration)',
            'tech': 'Python · ML · MongoDB · Docker · AWS · GridSearchCV',
            'bullets': [
                'Predicted US visa approval outcomes using production-grade ML models; performed EDA, feature engineering, and hyperparameter tuning via GridSearchCV.',
                'Connected MongoDB for data persistence; containerised solution with Docker for AWS deployment.',
            ],
            'relevant_to': ['data scientist', 'ml engineer'],
        },
        {
            'title': 'Disease Classification Mobile App',
            'tech': 'TensorFlow · CNN · TFLite · FastAPI · GCP · ReactJS · React Native',
            'bullets': [
                'Achieved 92% accuracy detecting Early/Late Blight in potato crops; quantized CNN with TF Lite — reduced model size by 80% and inference latency by 50% for on-device deployment.',
                'Deployed serverless inference on Google Cloud Functions (25% cost reduction); served real-time predictions via TF Serving + FastAPI to ReactJS and React Native frontends.',
            ],
            'relevant_to': ['ai engineer', 'ml engineer', 'data scientist'],
        },
        {
            'title': 'Accident & Fall Detection System',
            'tech': 'Python · YOLO · OpenCV · CNN · CUDA · SMTP',
            'bullets': [
                'Built real-time human fall detection system achieving 95% accuracy using YOLO + CUDA toolkit on edge hardware.',
                'Designed smart camera surveillance for highway accident detection; automated ambulance dispatch via SMTP alerts, cutting emergency response time by 40%.',
            ],
            'relevant_to': ['ai engineer', 'computer vision'],
        },
    ]
    
    # Filter projects: keep relevant ones to the job
    job_lower = job_title.lower()
    for proj in projects:
        # Keep project if any of its relevance tags match the job
        is_relevant = any(tag in job_lower for tag in proj['relevant_to'])
        # For Databricks AI Engineer, all ML/AI projects are relevant
        if 'ai engineer' in job_lower or 'data scientist' in job_lower:
            is_relevant = True
        
        if is_relevant:
            p = doc.add_paragraph()
            run = p.add_run(f"{proj['title']}  ")
            run.font.bold = True
            run.font.size = Pt(9.5)
            run = p.add_run(f"| {proj['tech']}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(4)
            
            for bullet in proj['bullets']:
                add_bullet(doc, bullet)
    
    # ═══════════════════════════════════════════════════════════════════
    # EDUCATION (exact same as original)
    # ═══════════════════════════════════════════════════════════════════
    
    add_section_heading(doc, 'EDUCATION')
    
    p = doc.add_paragraph()
    run = p.add_run('M.Tech in Data Science & Engineering')
    run.font.bold = True
    run.font.size = Pt(9.5)
    run = p.add_run(' | BITS Pilani, India | Oct 2021 – Oct 2023')
    run.font.size = Pt(9.5)
    p.paragraph_format.space_after = Pt(2)
    
    p = doc.add_paragraph()
    run = p.add_run('B.E in Computer Science & Engineering')
    run.font.bold = True
    run.font.size = Pt(9.5)
    run = p.add_run(' | OIST Bhopal, India | Aug 2015 – Jun 2019')
    run.font.size = Pt(9.5)
    
    # ═══════════════════════════════════════════════════════════════════
    # CERTIFICATIONS & AWARDS (exact same as original)
    # ═══════════════════════════════════════════════════════════════════
    
    add_section_heading(doc, 'CERTIFICATIONS & AWARDS')
    
    certs = [
        'GCP Professional ML Engineer (Certified Feb 2023)',
        'Certified MLOps Developer – Dataiku',
        'IBM Professional Data Science Certification',
        '1st Prize – Trizetto Hackathon (Patient Data Integration, 95% approval rate)',
        'OpenAI Whisper Hackathon – Hearing-impaired accessibility app',
        'Secretary, Toastmasters Club · Speaker, PyData Conference',
    ]
    
    for cert in certs:
        add_bullet(doc, cert)
    
    # ═══════════════════════════════════════════════════════════════════
    # Save
    # ═══════════════════════════════════════════════════════════════════
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M')
    company_clean = re.sub(r'[^a-zA-Z0-9]', '_', company)
    title_clean = re.sub(r'[^a-zA-Z0-9]', '_', job_title)[:40]
    
    # Create timestamped subfolder in artifacts
    run_folder = f"data/artifacts/{datetime.now().strftime('%Y%m%d_%H%M')}"
    os.makedirs(run_folder, exist_ok=True)
    
    # Save to generated-resumes
    os.makedirs("data/generated-resumes", exist_ok=True)
    filename = f"MohammadSheriff_{company_clean}_{title_clean}_{timestamp}.docx"
    filepath = f"data/generated-resumes/{filename}"
    doc.save(filepath)
    
    return filepath


def add_section_heading(doc, text):
    """Add a section heading with underline (matching original style)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(30, 30, 30)
    
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): '000000',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_company_header(doc, company_line, title, dates):
    """Add company/role header exactly like original."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(company_line)
    run.font.bold = True
    run.font.size = Pt(10)
    
    # Title and dates on same line
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(9.5)
    run = p.add_run(f'  |  {dates}')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(100, 100, 100)


def add_bullet(doc, text):
    """Add a bullet point with triangle marker like original (▸)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run('▸ ')
    run.font.size = Pt(9.5)
    run = p.add_run(text)
    run.font.size = Pt(9.5)


# ═══════════════════════════════════════════════════════════════════════
# Generate tailored resumes for top matches
# ═══════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    print("=" * 60)
    print("GENERATING TAILORED RESUMES (Following Original Format)")
    print("=" * 60)
    
    # Databricks AI Engineer - India (Remote)
    path1 = create_tailored_resume(
        job_title="AI Engineer - FDE (Forward Deployed Engineer)",
        company="Databricks",
        location="Remote - India",
        jd_keywords=["GenAI", "RAG", "multi-agent systems", "Text2SQL", "fine-tuning",
                     "HuggingFace", "LangChain", "DSPy", "PyTorch", "pandas", "scikit-learn",
                     "AWS", "Azure", "GCP", "production-grade", "LLMOps", "Mosaic AI",
                     "Data + AI Summit", "technical advisor", "cross-functional"],
    )
    print(f"\n✅ Databricks (India Remote): {path1}")
    
    # Databricks AI Engineer - Sydney  
    path2 = create_tailored_resume(
        job_title="AI Engineer - FDE (Forward Deployed Engineer)",
        company="Databricks",
        location="Sydney, Australia",
        jd_keywords=["GenAI", "RAG", "multi-agent systems", "Text2SQL", "fine-tuning",
                     "HuggingFace", "LangChain", "DSPy", "PyTorch", "pandas", "scikit-learn",
                     "AWS", "Azure", "GCP", "production-grade", "LLMOps"],
    )
    print(f"✅ Databricks (Sydney): {path2}")
    
    # Airbnb AI Engineer
    path3 = create_tailored_resume(
        job_title="AI Engineer, Community Support Engineering",
        company="Airbnb",
        location="China - Remote",
        jd_keywords=["AI", "Machine Learning", "RAG", "scalable services",
                     "cross-functional", "experimentation", "AI-powered applications"],
    )
    print(f"✅ Airbnb (Remote): {path3}")
    
    # Convert to PDF
    print("\n" + "=" * 60)
    print("CONVERTING TO PDF...")
    print("=" * 60)
    
    try:
        from docx2pdf import convert
        
        for docx_path in [path1, path2, path3]:
            pdf_path = docx_path.replace('.docx', '.pdf')
            convert(docx_path, pdf_path)
            print(f"  📄 {pdf_path}")
        
        print("\n✅ All PDFs generated successfully!")
    except Exception as e:
        print(f"  ⚠️ PDF conversion error: {e}")
        print("  DOCX files are still available.")
    
    print("\n" + "=" * 60)
    print("FILES READY:")
    print("=" * 60)
    print(f"  Original:  data/actual-resume/MohammadSheriff_FAANG_v5.pdf")
    print(f"  Tailored:  {path1}")
    print(f"  Tailored:  {path2}")
    print(f"  Tailored:  {path3}")
    print("\nAPPLY HERE:")
    print("  1. Databricks (India Remote): https://databricks.com/company/careers/open-positions/job?gh_jid=8099751002")
    print("  2. Databricks (Sydney):       https://databricks.com/company/careers/open-positions/job?gh_jid=8298792002")
    print("  3. Airbnb (Remote):           https://careers.airbnb.com/positions/7946288?gh_jid=7946288")
