"""Resume DOCX formatter - Builds ATS-optimized 2-page resume with hyperlinks and consistent formatting.

This formatter follows the exact template from the user's original resume:
- Page 1: Header → Summary → Skills → ValueLabs → LTIMindtree
- Page 2: Cognizant → Research Projects → Education → Certifications
- Consistent · separator throughout
- ▸ bullet markers
- Clickable hyperlinks (LinkedIn, GitHub, Portfolio, PyPI, project links)
- 4-5 bullets per company with metrics
"""

import json
import re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.resume.tailoring_agent import TailoringResult, TailoredSection


# Resume data extracted from actual resume (source of truth)
RESUME_DATA = {
    "name": "Mohammad Sheriff Mehmood",
    "title": "Senior Data Scientist · NLP & Generative AI Engineer",
    "phone": "+91 7999937693",
    "email": "mdsheriff2702@gmail.com",
    "linkedin_url": "https://www.linkedin.com/in/mohammad-sheriff/",
    "github_url": "https://github.com/sheriff786",
    "portfolio_url": "https://sheriff786.github.io/My_Portfolio_Website/#about",
    "authentica_url": "https://pypi.org/project/authentica/",
    "skills": {
        "Languages": "Python (primary) — proficient in TensorFlow · PyTorch · Scikit-learn · NumPy · Pandas · OpenCV · YOLO · CNN",
        "GenAI & LLMs": "LangChain · LangSmith · OpenAI API · Gemini · RAG · LLMOps · Fine-tuning · Prompt Engineering · Agentic AI · MCP · Transformer Models · Sentence-Transformers · CrewAI",
        "MLOps & Cloud": "MLflow · Docker · Kubernetes · Airflow · GCP (Vertex AI) · AWS (SageMaker · ECS · CodePipeline) · CI/CD",
        "Databases": "PostgreSQL · MongoDB · SQL · ChromaDB · FAISS · Pinecone",
        "Tools & Other": "Tableau · Weights & Biases · FastAPI · Flask · Django · A/B Testing · REST API Design · Git",
    },
    "experience": [
        {
            "company": "ValueLabs",
            "location": "Hyderabad, India",
            "title": "Senior Data Scientist | AI/ML Engineer",
            "dates": "Mar 2025 – Present",
            "context": "Embedded with Businessolver's Data Science team as a full-stack ML Engineer; designed and shipped an automated document analysis pipeline leveraging computer vision · VLMs · and AWS-native tooling to process healthcare documents at scale.",
            "bullets": [
                "Document Analysis Pipeline: Designed an end-to-end multi-modal pipeline on AWS (Comprehend · Textract · S3 · ECS) to classify and extract data from 47,000+ healthcare documents · significantly cutting manual claim validation time.",
                "Query-Based Document Understanding: Built a custom QA system integrating visual features from Textract OCR with textual embeddings across varied layouts; deployed inference on AWS Lambda achieving 93% extraction precision.",
                "Annotation & Monitoring UI: Built an interactive React UI with REST API for data annotation/validation · real-time model metric visualisation · performance monitoring · and data-driven threshold tuning streamlining stakeholder review cycles.",
                "Stacking Ensemble Model: Trained and tuned a stacking ensemble (Random Forest · XGBoost · LightGBM) via cross-validation achieving 0.78 F1 and 0.81 AUC; applied SHAP and Gini impurity for model explainability.",
            ],
        },
        {
            "company": "LTIMindtree",
            "location": "Pune, India",
            "title": "Senior Data Scientist – NLP Engineer",
            "dates": "Jan 2024 – Mar 2025",
            "context": None,
            "bullets": [
                "GenAI Chatbot & RAG System: Architected production-grade RAG chatbot (Gemini + FAISS Vector DB) serving 15+ team members with 90% query accuracy and sub-2s response time for enterprise knowledge retrieval.",
                "LLMOps Pipeline: Built end-to-end LLMOps pipeline using OpenAI · LangChain · ChromaDB · and sentence-transformers cutting model iteration time by 60% and reducing deployment overhead by 50%.",
                "Healthcare Revenue Model: Deployed anomaly-detection targeting model generating $1M+ in annual healthcare revenue at 90% precision; reduced false positives by 20% via class-weighted ensemble (Logistic Regression + Random Forest).",
                "Workflow Orchestration: Designed Airflow DAGs to automate multi-step ML pipelines · cutting end-to-end processing time by 40% and improving delivery velocity across a 7-member team.",
                "Semantic NLP (BioBERT): Generated vector embeddings for ICD-9 and medical procedure codes using BioBERT · improving semantic search accuracy by 18% over TF-IDF baseline.",
            ],
        },
        {
            "company": "Cognizant Technology Solutions",
            "location": "Chennai, India",
            "title": "Data Scientist",
            "dates": "Dec 2019 – Jan 2024",
            "context": "Joined as Programmer Analyst → core ML systems, promoted to Data Scientist (leading cross-functional AI/NLP initiatives).",
            "page_break_before": True,
            "bullets": [
                "ML Model Development: Built and deployed supervised ML models improving prediction accuracy by 20% and saving $200K in operational costs; improved baselines by 5–10% using Neural Networks · ensemble methods · GridSearchCV.",
                "AI Customer Support Platform: Led 7-member cross-functional team to build NLP-powered support application; reduced customer churn by 40% and generated $300K additional quarterly revenue through context-aware intelligent responses.",
                "Fraud Detection System: Reduced false-positive fraud alerts by 20% using Logistic Regression and Random Forest with class-weighted resampling — received company Innovation Award.",
                "Data Platform & BI: Unified 10+ heterogeneous data sources (databases · APIs · third-party feeds) into ML-ready pipelines; built Tableau dashboards tracking 20+ KPIs · improving cross-team reporting efficiency by 30%.",
            ],
        },
    ],
    "projects": [
        {
            "title": "Abstractive Text Summarisation",
            "tech": "Python · BERT · NLP · Transformers",
            "github_url": "https://github.com/sheriff786",
            "bullets": [
                "Applied BERT and NLP toolkits to summarise live chat context, improving agent efficiency by 35% and reducing customer response times by 25%.",
                "Gave agents instant access to previous chat history — reducing handle time and improving resolution quality.",
            ],
        },
        {
            "title": "End-to-End Topic Modelling & MLOps on AWS",
            "tech": "Python · LDA · Docker · AWS ECS · MLflow · CodePipeline · EC2",
            "github_url": "https://github.com/sheriff786",
            "bullets": [
                "Built LDA topic modelling pipeline on AWS ECS with blue/green deployment — cut processing time by 40% and deployment time by 60%.",
                "Automated full CI/CD (CodeCommit · CodeBuild · CodeDeploy · CodePipeline); reduced manual effort by 70%. Used MLflow for model versioning.",
            ],
        },
        {
            "title": "Visa Approval Forecasting (US Immigration)",
            "tech": "Python · ML · MongoDB · Docker · AWS · GridSearchCV",
            "github_url": "https://github.com/sheriff786",
            "bullets": [
                "Predicted US visa approval outcomes using production-grade ML models; performed EDA · feature engineering · and hyperparameter tuning via GridSearchCV.",
                "Connected MongoDB for data persistence; containerised solution with Docker for AWS deployment.",
            ],
        },
        {
            "title": "Disease Classification Mobile App",
            "tech": "TensorFlow · CNN · TFLite · FastAPI · GCP · ReactJS · React Native",
            "github_url": "https://github.com/sheriff786",
            "bullets": [
                "Achieved 92% accuracy detecting Early/Late Blight; quantized CNN with TF Lite — reduced model size by 80% and inference latency by 50% for on-device deployment.",
                "Deployed serverless inference on Google Cloud Functions (25% cost reduction); served real-time predictions via TF Serving + FastAPI.",
            ],
        },
        {
            "title": "Accident & Fall Detection System",
            "tech": "Python · YOLO · OpenCV · CNN · CUDA · SMTP",
            "github_url": "https://github.com/sheriff786",
            "bullets": [
                "Built real-time human fall detection system achieving 95% accuracy using YOLO + CUDA toolkit on edge hardware.",
                "Designed smart camera surveillance for highway accident detection; automated ambulance dispatch via SMTP alerts · cutting emergency response time by 40%.",
            ],
        },
    ],
    "education": [
        {"degree": "M.Tech in Data Science & Engineering", "institution": "BITS Pilani, India", "dates": "Oct 2021 – Oct 2023"},
        {"degree": "B.E in Computer Science & Engineering", "institution": "OIST Bhopal, India", "dates": "Aug 2015 – Jun 2019"},
    ],
    "certifications": [
        "GCP Professional ML Engineer (Certified Feb 2023)",
        "Certified MLOps Developer – Dataiku",
        "IBM Professional Data Science Certification",
        "1st Prize – Trizetto Hackathon (Patient Data Integration · 95% approval rate)",
        "OpenAI Whisper Hackathon – Hearing-impaired accessibility app",
        "Secretary · Toastmasters Club  ·  Speaker · PyData Conference",
    ],
}


class ResumeFormatter:
    """Builds ATS-optimized 2-page DOCX resume with hyperlinks and consistent formatting."""

    def __init__(self, template_path: str = None):
        self.output_dir = Path("data/generated-resumes")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.resume_data = RESUME_DATA

    def read_sections(self) -> dict[str, str]:
        """Return resume sections as text for the tailoring agent."""
        sections = {
            "header": f"{self.resume_data['name']}\n{self.resume_data['title']}",
            "professional summary": self._get_base_summary(),
            "technical skills": "\n".join(f"{k}: {v}" for k, v in self.resume_data["skills"].items()),
            "professional experience": self._get_experience_text(),
            "research projects": self._get_projects_text(),
            "education": "\n".join(f"{e['degree']} · {e['institution']} · {e['dates']}" for e in self.resume_data["education"]),
            "certifications & awards": "\n".join(self.resume_data["certifications"]),
        }
        return sections

    def write_tailored_resume(self, tailoring_result: TailoringResult) -> str:
        """Write a new tailored DOCX using the structured approach."""
        tailored_map = {s.section_name.lower(): s for s in tailoring_result.sections_modified}

        summary = None
        for key in ("summary", "professional summary", "profile"):
            if key in tailored_map:
                summary = tailored_map[key].tailored_content
                break

        doc = self._build_document(tailored_summary=summary, company=tailoring_result.company)
        output_path = tailoring_result.output_path
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return output_path

    def build_tailored_resume(self, company: str, title: str, tailored_summary: str = None, extra_skill_keywords: list[str] = None) -> str:
        """Build a tailored resume DOCX directly."""
        doc = self._build_document(tailored_summary=tailored_summary, extra_skill_keywords=extra_skill_keywords, company=company)
        safe_company = re.sub(r"[^\w\-]", "_", company)
        safe_title = re.sub(r"[^\w\-]", "_", title)[:40]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        filename = f"MohammadSheriff_{safe_company}_{safe_title}_{timestamp}"
        docx_path = self.output_dir / f"{filename}.docx"
        doc.save(str(docx_path))
        return str(docx_path)

    def _build_document(self, tailored_summary=None, extra_skill_keywords=None, company="") -> Document:
        """Build complete DOCX following exact original format."""
        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(9.5)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)

        self._add_header(doc)
        self._add_section_heading(doc, "PROFESSIONAL SUMMARY")
        self._add_summary(doc, tailored_summary)
        self._add_section_heading(doc, "TECHNICAL SKILLS")
        self._add_skills(doc, extra_skill_keywords)
        self._add_section_heading(doc, "PROFESSIONAL EXPERIENCE")
        for exp in self.resume_data["experience"]:
            if exp.get("page_break_before"):
                p = doc.add_paragraph()
                p.add_run().add_break(WD_BREAK.PAGE)
            self._add_experience_entry(doc, exp)
        self._add_section_heading(doc, "RESEARCH PROJECTS")
        for proj in self.resume_data["projects"]:
            self._add_project_entry(doc, proj)
        self._add_section_heading(doc, "EDUCATION")
        for edu in self.resume_data["education"]:
            p = doc.add_paragraph()
            self._set_spacing(p, before=3, after=2)
            run = p.add_run(edu["degree"])
            run.font.bold = True
            run.font.size = Pt(9.5)
            p.add_run(f"  ·  {edu['institution']}  ·  {edu['dates']}").font.size = Pt(9.5)
        self._add_section_heading(doc, "CERTIFICATIONS & AWARDS")
        for cert in self.resume_data["certifications"]:
            self._add_bullet(doc, cert)

        return doc

    def _add_header(self, doc):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_spacing(p, after=1)
        run = p.add_run(self.resume_data["name"])
        run.font.size = Pt(16)
        run.font.bold = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_spacing(p, after=2)
        run = p.add_run(self.resume_data["title"])
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(60, 60, 60)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_spacing(p, after=6)
        p.add_run(f"{self.resume_data['phone']} · {self.resume_data['email']} · ").font.size = Pt(9)
        self._add_hyperlink(p, "LinkedIn", self.resume_data["linkedin_url"], Pt(9))
        p.add_run(" · ").font.size = Pt(9)
        self._add_hyperlink(p, "github.com/sheriff786", self.resume_data["github_url"], Pt(9))
        p.add_run(" · ").font.size = Pt(9)
        self._add_hyperlink(p, "Portfolio", self.resume_data["portfolio_url"], Pt(9))

    def _add_summary(self, doc, tailored_summary=None):
        text = tailored_summary or self._get_base_summary()
        p = doc.add_paragraph()
        self._set_spacing(p, before=3, after=4)
        if "authentica" in text:
            parts = text.split("authentica", 1)
            p.add_run(parts[0]).font.size = Pt(9.5)
            self._add_hyperlink(p, "authentica", self.resume_data["authentica_url"], Pt(9.5))
            p.add_run(parts[1]).font.size = Pt(9.5)
        else:
            p.add_run(text).font.size = Pt(9.5)

    def _add_skills(self, doc, extra_keywords=None):
        skills = dict(self.resume_data["skills"])
        if extra_keywords:
            genai = skills.get("GenAI & LLMs", "")
            for kw in extra_keywords:
                if kw.lower() not in genai.lower():
                    genai += f" · {kw}"
            skills["GenAI & LLMs"] = genai

        for category, items in skills.items():
            p = doc.add_paragraph()
            self._set_spacing(p, before=1, after=1)
            run = p.add_run(f"{category}:  ")
            run.font.bold = True
            run.font.size = Pt(9.5)
            p.add_run(items).font.size = Pt(9.5)

    def _add_experience_entry(self, doc, exp):
        p = doc.add_paragraph()
        self._set_spacing(p, before=6 if not exp.get("page_break_before") else 0, after=0)
        run = p.add_run(f"{exp['company']} · {exp['location']}")
        run.font.bold = True
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        self._set_spacing(p, before=1, after=0)
        run = p.add_run(exp["title"])
        run.font.bold = True
        run.font.size = Pt(9.5)

        p = doc.add_paragraph()
        self._set_spacing(p, before=0, after=2)
        run = p.add_run(exp["dates"])
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(80, 80, 80)

        if exp.get("context"):
            p = doc.add_paragraph()
            self._set_spacing(p, before=1, after=2)
            run = p.add_run(exp["context"])
            run.font.italic = True
            run.font.size = Pt(9.5)

        for bullet in exp["bullets"]:
            self._add_bullet(doc, bullet)

    def _add_project_entry(self, doc, proj):
        p = doc.add_paragraph()
        self._set_spacing(p, before=4, after=1)
        run = p.add_run(proj["title"])
        run.font.bold = True
        run.font.size = Pt(9.5)
        run = p.add_run(f"  |  {proj['tech']}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)
        if proj.get("github_url"):
            p.add_run("    ").font.size = Pt(8.5)
            self._add_hyperlink(p, "View on GitHub →", proj["github_url"], Pt(8.5))
        for bullet in proj["bullets"]:
            self._add_bullet(doc, bullet)

    def _add_section_heading(self, doc, text):
        p = doc.add_paragraph()
        self._set_spacing(p, before=8, after=3)
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0, 0, 0)
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "4",
            qn("w:space"): "1", qn("w:color"): "333333",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_bullet(self, doc, text, indent=0.4):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent)
        self._set_spacing(p, before=1, after=1)
        p.add_run("▸ ").font.size = Pt(9.5)
        p.add_run(text).font.size = Pt(9.5)

    def _add_hyperlink(self, paragraph, text, url, font_size=Pt(9)):
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        c = OxmlElement("w:color")
        c.set(qn("w:val"), "0066CC")
        rPr.append(c)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size.pt * 2)))
        rPr.append(sz)
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _set_spacing(self, p, before=0, after=0):
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)

    def _get_base_summary(self):
        return (
            "Senior Data Scientist & AI/ML Engineer with 6+ years of experience. "
            "I build AI systems that move business metrics — not just demos. "
            "Specializing in NLP · LLMs · and production MLOps, I have shipped enterprise-grade "
            "GenAI systems that delivered $1M+ revenue uplift · cut churn by 40% · and accelerated "
            "model iteration by 60% across healthcare · finance · and retail. "
            "Open-source author: published authentica on PyPI, actively used in production "
            "at Businessolver (Healthcare Benefits SaaS)."
        )

    def _get_experience_text(self):
        lines = []
        for exp in self.resume_data["experience"]:
            lines.append(f"{exp['company']} · {exp['location']}")
            lines.append(f"{exp['title']} | {exp['dates']}")
            if exp.get("context"):
                lines.append(exp["context"])
            for b in exp["bullets"]:
                lines.append(f"▸ {b}")
            lines.append("")
        return "\n".join(lines)

    def _get_projects_text(self):
        lines = []
        for proj in self.resume_data["projects"]:
            lines.append(f"{proj['title']} | {proj['tech']}")
            for b in proj["bullets"]:
                lines.append(f"▸ {b}")
            lines.append("")
        return "\n".join(lines)
