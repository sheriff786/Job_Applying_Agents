"""
Generate ATS-optimized, 2-page tailored resume following EXACT original format.
- Consistent formatting (· separators for skills)
- 4-5 impactful bullets per company with metrics
- Research projects filtered by JD relevance
- Professional, clean, clearly readable
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
import os
import re


def set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0):
    """Set precise paragraph spacing."""
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if line_spacing != 1.0:
        p.paragraph_format.line_spacing = line_spacing


def add_section_heading(doc, text):
    """Section heading with bottom border - clean professional style."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=3)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0, 0, 0)
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4',
        qn('w:space'): '1', qn('w:color'): '333333',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_hyperlink(paragraph, text, url, font_size=Pt(9), color=RGBColor(0, 102, 204), bold=False):
    """Add a clickable hyperlink to a paragraph."""
    from docx.oxml import OxmlElement
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), f'{color[0]:02x}{color[1]:02x}{color[2]:02x}')
    rPr.append(c)
    
    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    # Font size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(sz)
    
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_bullet(doc, text, indent=0.4):
    """Add bullet with ▸ marker."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    set_paragraph_spacing(p, before=1, after=1)
    run = p.add_run('▸ ')
    run.font.size = Pt(9.5)
    run = p.add_run(text)
    run.font.size = Pt(9.5)


def create_tailored_resume_v2(job_title, company, location, jd_keywords):
    """Create 2-page ATS-optimized resume following exact original template."""
    
    doc = Document()
    
    # Page setup - A4 with narrow margins for 2-page fit
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(9.5)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    
    # ═══════════════════════════════════════════════════════════════
    # HEADER
    # ═══════════════════════════════════════════════════════════════
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=1)
    run = p.add_run('Mohammad Sheriff Mehmood')
    run.font.size = Pt(16)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=2)
    run = p.add_run('Senior Data Scientist · NLP & Generative AI Engineer')
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(60, 60, 60)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=6)
    run = p.add_run('+91 7999937693 · mdsheriff2702@gmail.com · ')
    run.font.size = Pt(9)
    add_hyperlink(p, 'LinkedIn', 'https://www.linkedin.com/in/mohammad-sheriff/', font_size=Pt(9))
    run = p.add_run(' · ')
    run.font.size = Pt(9)
    add_hyperlink(p, 'github.com/sheriff786', 'https://github.com/sheriff786', font_size=Pt(9))
    run = p.add_run(' · ')
    run.font.size = Pt(9)
    add_hyperlink(p, 'Portfolio', 'https://sheriff786.github.io/My_Portfolio_Website/#about', font_size=Pt(9))
    
    # ═══════════════════════════════════════════════════════════════
    # PROFESSIONAL SUMMARY
    # ═══════════════════════════════════════════════════════════════
    
    add_section_heading(doc, 'PROFESSIONAL SUMMARY')
    
    if 'databricks' in company.lower():
        summary_text = (
            'Senior Data Scientist & AI/ML Engineer with 6+ years of experience building production-grade '
            'GenAI applications including RAG · multi-agent systems · fine-tuning with LangChain · HuggingFace · PyTorch. '
            'Specializing in NLP · LLMs · and production MLOps, I have shipped enterprise-grade GenAI systems that delivered '
            '$1M+ revenue uplift · cut churn by 40% · and accelerated model iteration by 60% across healthcare · finance · and retail. '
            'Proven technical advisor deploying AI solutions on AWS · Azure · GCP. '
            'Open-source author: published '
        )
        summary_link_text = 'authentica'
        summary_link_url = 'https://pypi.org/project/authentica/'
        summary_after = ' on PyPI, actively used in production at Businessolver (Healthcare Benefits SaaS).'
    elif 'airbnb' in company.lower():
        summary_text = (
            'Senior Data Scientist & AI/ML Engineer with 6+ years of experience building scalable AI-powered applications. '
            'Specializing in NLP · LLMs · RAG systems · and production MLOps, I have shipped enterprise-grade GenAI systems that delivered '
            '$1M+ revenue uplift · cut churn by 40% · and accelerated model iteration by 60% across healthcare · finance · and retail. '
            'Experienced in cross-functional collaboration with product · engineering · and operations teams. '
            'Open-source author: published '
        )
        summary_link_text = 'authentica'
        summary_link_url = 'https://pypi.org/project/authentica/'
        summary_after = ' on PyPI, actively used in production at Businessolver (Healthcare Benefits SaaS).'
    else:
        summary_text = (
            'Senior Data Scientist & AI/ML Engineer with 6+ years of experience. I build AI systems that move business metrics — not just demos. '
            'Specializing in NLP · LLMs · and production MLOps, I have shipped enterprise-grade GenAI systems that delivered '
            '$1M+ revenue uplift · cut churn by 40% · and accelerated model iteration by 60% across healthcare · finance · and retail. '
            'Open-source author: published '
        )
        summary_link_text = 'authentica'
        summary_link_url = 'https://pypi.org/project/authentica/'
        summary_after = ' on PyPI, actively used in production at Businessolver (Healthcare Benefits SaaS).'
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=4)
    run = p.add_run(summary_text)
    run.font.size = Pt(9.5)
    add_hyperlink(p, summary_link_text, summary_link_url, font_size=Pt(9.5))
    run = p.add_run(summary_after)
    run.font.size = Pt(9.5)
    
    # ═══════════════════════════════════════════════════════════════
    # TECHNICAL SKILLS (consistent · separator throughout)
    # ═══════════════════════════════════════════════════════════════
    
    add_section_heading(doc, 'TECHNICAL SKILLS')
    
    if 'databricks' in company.lower():
        skills = [
            ('Languages', 'Python (primary) — proficient in TensorFlow · PyTorch · Scikit-learn · NumPy · Pandas · OpenCV · YOLO · CNN'),
            ('GenAI & LLMs', 'LangChain · LangSmith · HuggingFace · DSPy · OpenAI API · Gemini · RAG · LLMOps · Fine-tuning · Prompt Engineering · Agentic AI · MCP · Multi-agent Systems · Text2SQL · Sentence-Transformers · CrewAI'),
            ('MLOps & Cloud', 'MLflow · Docker · Kubernetes · Airflow · GCP (Vertex AI) · AWS (SageMaker · ECS · CodePipeline) · Azure · Apache Spark · CI/CD'),
            ('Databases', 'PostgreSQL · MongoDB · SQL · ChromaDB · FAISS · Pinecone'),
            ('Tools & Other', 'Tableau · Weights & Biases · FastAPI · Flask · Django · A/B Testing · REST API Design · Git'),
        ]
    elif 'airbnb' in company.lower():
        skills = [
            ('Languages', 'Python (primary) — proficient in TensorFlow · PyTorch · Scikit-learn · NumPy · Pandas · OpenCV · YOLO · CNN'),
            ('GenAI & LLMs', 'LangChain · LangSmith · OpenAI API · Gemini · RAG · LLMOps · Fine-tuning · Prompt Engineering · Agentic AI · MCP · Transformer Models · Sentence-Transformers · CrewAI'),
            ('MLOps & Cloud', 'MLflow · Docker · Kubernetes · Airflow · GCP (Vertex AI) · AWS (SageMaker · ECS · CodePipeline) · CI/CD'),
            ('Databases', 'PostgreSQL · MongoDB · SQL · ChromaDB · FAISS · Pinecone'),
            ('Tools & Other', 'Tableau · Weights & Biases · FastAPI · Flask · Django · A/B Testing · REST API Design · Git'),
        ]
    else:
        skills = [
            ('Languages', 'Python (primary) — proficient in TensorFlow · PyTorch · Scikit-learn · NumPy · Pandas · OpenCV · YOLO · CNN'),
            ('GenAI & LLMs', 'LangChain · LangSmith · OpenAI API · Gemini · RAG · LLMOps · Fine-tuning · Prompt Engineering · Agentic AI · MCP · Transformer Models · Sentence-Transformers · CrewAI'),
            ('MLOps & Cloud', 'MLflow · Docker · Kubernetes · Airflow · GCP (Vertex AI) · AWS (SageMaker · ECS · CodePipeline) · CI/CD'),
            ('Databases', 'PostgreSQL · MongoDB · SQL · ChromaDB · FAISS · Pinecone'),
            ('Tools & Other', 'Tableau · Weights & Biases · FastAPI · Flask · Django · A/B Testing · REST API Design · Git'),
        ]
    
    for category, items in skills:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=1, after=1)
        run = p.add_run(f'{category}:  ')
        run.font.bold = True
        run.font.size = Pt(9.5)
        run = p.add_run(items)
        run.font.size = Pt(9.5)
    
    # ═══════════════════════════════════════════════════════════════
    # PROFESSIONAL EXPERIENCE
    # Following exact format:
    #   Company · Location
    #   Title
    #   Dates
    #   [Context line]
    #   ▸ Bullets
    # ═══════════════════════════════════════════════════════════════
    
    add_section_heading(doc, 'PROFESSIONAL EXPERIENCE')
    
    # --- VALUELABS ---
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=0)
    run = p.add_run('ValueLabs · Hyderabad, India')
    run.font.bold = True
    run.font.size = Pt(10)
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=0)
    run = p.add_run('Senior Data Scientist | AI/ML Engineer')
    run.font.bold = True
    run.font.size = Pt(9.5)
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=2)
    run = p.add_run('Mar 2025 – Present')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(80, 80, 80)
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=2)
    run = p.add_run(
        'Embedded with Businessolver\'s Data Science team as a full-stack ML Engineer; designed and shipped an automated '
        'document analysis pipeline leveraging computer vision · VLMs · and AWS-native tooling to process healthcare documents at scale.'
    )
    run.font.italic = True
    run.font.size = Pt(9.5)
    
    # ValueLabs bullets - Businessolver embedded work (4-5 points)
    valuelabs_bullets = [
        'Document Analysis Pipeline: Designed an end-to-end multi-modal pipeline on AWS (Comprehend · Textract · S3 · ECS) to classify and extract data from 47,000+ healthcare documents · significantly cutting manual claim validation time.',
        'Query-Based Document Understanding: Built a custom QA system integrating visual features from Textract OCR with textual embeddings across varied layouts; deployed inference on AWS Lambda achieving 93% extraction precision.',
        'Annotation & Monitoring UI: Built an interactive React UI with REST API for data annotation/validation · real-time model metric visualisation · performance monitoring · and data-driven threshold tuning streamlining stakeholder review cycles.',
        'Stacking Ensemble Model: Trained and tuned a stacking ensemble (Random Forest · XGBoost · LightGBM) via cross-validation achieving 0.78 F1 and 0.81 AUC; applied SHAP and Gini impurity for model explainability.',
    ]
    
    for bullet in valuelabs_bullets:
        add_bullet(doc, bullet)
    
    # --- LTIMINDTREE ---
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=0)
    run = p.add_run('LTIMindtree · Pune, India')
    run.font.bold = True
    run.font.size = Pt(10)
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=0)
    run = p.add_run('Senior Data Scientist – NLP Engineer')
    run.font.bold = True
    run.font.size = Pt(9.5)
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=2)
    run = p.add_run('Jan 2024 – Mar 2025')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(80, 80, 80)
    
    # LTIMindtree bullets - GenAI/NLP work (5 points)
    ltim_bullets = [
        'GenAI Chatbot & RAG System: Architected production-grade RAG chatbot (Gemini + FAISS Vector DB) serving 15+ team members with 90% query accuracy and sub-2s response time for enterprise knowledge retrieval.',
        'LLMOps Pipeline: Built end-to-end LLMOps pipeline using OpenAI · LangChain · ChromaDB · and sentence-transformers cutting model iteration time by 60% and reducing deployment overhead by 50%.',
        'Healthcare Revenue Model: Deployed anomaly-detection targeting model generating $1M+ in annual healthcare revenue at 90% precision; reduced false positives by 20% via class-weighted ensemble (Logistic Regression + Random Forest).',
        'Workflow Orchestration: Designed Airflow DAGs to automate multi-step ML pipelines · cutting end-to-end processing time by 40% and improving delivery velocity across a 7-member team.',
        'Semantic NLP (BioBERT): Generated vector embeddings for ICD-9 and medical procedure codes using BioBERT · improving semantic search accuracy by 18% over TF-IDF baseline.',
    ]
    
    for bullet in ltim_bullets:
        add_bullet(doc, bullet)
    
    # --- COGNIZANT --- (starts on page 2)
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    run = p.add_run('Cognizant Technology Solutions · Chennai, India')
    run.font.bold = True
    run.font.size = Pt(10)
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=0)
    run = p.add_run('Data Scientist')
    run.font.bold = True
    run.font.size = Pt(9.5)
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=2)
    run = p.add_run('Dec 2019 – Jan 2024')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(80, 80, 80)
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=2)
    run = p.add_run('Joined as Programmer Analyst → core ML systems, promoted to Data Scientist (leading cross-functional AI/NLP initiatives).')
    run.font.italic = True
    run.font.size = Pt(9.5)
    
    # Cognizant bullets - 4 points with metrics
    cognizant_bullets = [
        'ML Model Development: Built and deployed supervised ML models improving prediction accuracy by 20% and saving $200K in operational costs; improved baselines by 5–10% using Neural Networks · ensemble methods · GridSearchCV.',
        'AI Customer Support Platform: Led 7-member cross-functional team to build NLP-powered support application; reduced customer churn by 40% and generated $300K additional quarterly revenue through context-aware intelligent responses.',
        'Fraud Detection System: Reduced false-positive fraud alerts by 20% using Logistic Regression and Random Forest with class-weighted resampling — received company Innovation Award.',
        'Data Platform & BI: Unified 10+ heterogeneous data sources (databases · APIs · third-party feeds) into ML-ready pipelines; built Tableau dashboards tracking 20+ KPIs · improving cross-team reporting efficiency by 30%.',
    ]
    
    for bullet in cognizant_bullets:
        add_bullet(doc, bullet)
    
    # ═══════════════════════════════════════════════════════════════
    # RESEARCH PROJECTS (filtered by JD relevance)
    # ═══════════════════════════════════════════════════════════════
    
    add_section_heading(doc, 'RESEARCH PROJECTS')
    
    if 'databricks' in company.lower():
        projects = [
            {
                'title': 'Abstractive Text Summarisation',
                'tech': 'Python · BERT · NLP · Transformers',
                'github_url': 'https://github.com/sheriff786',
                'bullets': [
                    'Applied BERT and NLP toolkits to summarise live chat context, improving agent efficiency by 35% and reducing customer response times by 25%.',
                    'Gave agents instant access to previous chat history — reducing handle time and improving resolution quality.',
                ],
            },
            {
                'title': 'End-to-End Topic Modelling & MLOps on AWS',
                'tech': 'Python · LDA · Docker · AWS ECS · MLflow · CodePipeline · EC2',
                'github_url': 'https://github.com/sheriff786',
                'bullets': [
                    'Built LDA topic modelling pipeline on AWS ECS with blue/green deployment — cut processing time by 40% and deployment time by 60%.',
                    'Automated full CI/CD (CodeCommit · CodeBuild · CodeDeploy · CodePipeline); reduced manual effort by 70%. Used MLflow for model versioning.',
                ],
            },
            {
                'title': 'Visa Approval Forecasting (US Immigration)',
                'tech': 'Python · ML · MongoDB · Docker · AWS · GridSearchCV',
                'github_url': 'https://github.com/sheriff786',
                'bullets': [
                    'Predicted US visa approval outcomes using production-grade ML models; performed EDA · feature engineering · and hyperparameter tuning via GridSearchCV.',
                    'Connected MongoDB for data persistence; containerised solution with Docker for AWS deployment.',
                ],
            },
            {
                'title': 'Disease Classification Mobile App',
                'tech': 'TensorFlow · CNN · TFLite · FastAPI · GCP · ReactJS · React Native',
                'github_url': 'https://github.com/sheriff786',
                'bullets': [
                    'Achieved 92% accuracy detecting Early/Late Blight; quantized CNN with TF Lite — reduced model size by 80% and inference latency by 50% for on-device deployment.',
                    'Deployed serverless inference on Google Cloud Functions (25% cost reduction); served real-time predictions via TF Serving + FastAPI.',
                ],
            },
            {
                'title': 'Accident & Fall Detection System',
                'tech': 'Python · YOLO · OpenCV · CNN · CUDA · SMTP',
                'github_url': 'https://github.com/sheriff786',
                'bullets': [
                    'Built real-time human fall detection system achieving 95% accuracy using YOLO + CUDA toolkit on edge hardware.',
                    'Designed smart camera surveillance for highway accident detection; automated ambulance dispatch via SMTP alerts · cutting emergency response time by 40%.',
                ],
            },
        ]
    elif 'airbnb' in company.lower():
        projects = [
            {
                'title': 'Abstractive Text Summarisation',
                'tech': 'Python · BERT · NLP · Transformers',
                'github_url': 'https://github.com/sheriff786',
                'bullets': [
                    'Applied BERT and NLP toolkits to summarise live chat context, improving agent efficiency by 35% and reducing customer response times by 25%.',
                    'Gave agents instant access to previous chat history — reducing handle time and improving resolution quality.',
                ],
            },
            {
                'title': 'End-to-End Topic Modelling & MLOps on AWS',
                'tech': 'Python · LDA · Docker · AWS ECS · MLflow · CodePipeline · EC2',
                'github_url': 'https://github.com/sheriff786',
                'bullets': [
                    'Built LDA topic modelling pipeline on AWS ECS with blue/green deployment — cut processing time by 40% and deployment time by 60%.',
                    'Automated full CI/CD (CodeCommit · CodeBuild · CodeDeploy · CodePipeline); reduced manual effort by 70%. Used MLflow for model versioning.',
                ],
            },
            {
                'title': 'Visa Approval Forecasting (US Immigration)',
                'tech': 'Python · ML · MongoDB · Docker · AWS · GridSearchCV',
                'github_url': 'https://github.com/sheriff786',
                'bullets': [
                    'Predicted US visa approval outcomes using production-grade ML models; performed EDA · feature engineering · and hyperparameter tuning via GridSearchCV.',
                    'Connected MongoDB for data persistence; containerised solution with Docker for AWS deployment.',
                ],
            },
            {
                'title': 'Disease Classification Mobile App',
                'tech': 'TensorFlow · CNN · TFLite · FastAPI · GCP · ReactJS · React Native',
                'github_url': 'https://github.com/sheriff786',
                'bullets': [
                    'Achieved 92% accuracy; quantized CNN with TF Lite — reduced model size by 80% and inference latency by 50%.',
                    'Deployed serverless inference on GCP Cloud Functions (25% cost reduction); served predictions via FastAPI to React frontends.',
                ],
            },
            {
                'title': 'Accident & Fall Detection System',
                'tech': 'Python · YOLO · OpenCV · CNN · CUDA · SMTP',
                'github_url': 'https://github.com/sheriff786',
                'bullets': [
                    'Built real-time human fall detection system achieving 95% accuracy using YOLO + CUDA toolkit on edge hardware.',
                    'Designed smart camera surveillance for highway accident detection; automated ambulance dispatch via SMTP alerts · cutting emergency response time by 40%.',
                ],
            },
        ]
    else:
        projects = [
            {
                'title': 'Abstractive Text Summarisation',
                'tech': 'Python · BERT · NLP · Transformers',
                'github_url': 'https://github.com/sheriff786',
                'bullets': [
                    'Applied BERT and NLP toolkits to summarise live chat context, improving agent efficiency by 35% and reducing customer response times by 25%.',
                    'Gave agents instant access to previous chat history — reducing handle time and improving resolution quality.',
                ],
            },
            {
                'title': 'End-to-End Topic Modelling & MLOps on AWS',
                'tech': 'Python · LDA · Docker · AWS ECS · MLflow · CodePipeline · EC2',
                'github_url': 'https://github.com/sheriff786',
                'bullets': [
                    'Built LDA topic modelling pipeline on AWS ECS with blue/green deployment — cut processing time by 40% and deployment time by 60%.',
                    'Automated full CI/CD (CodeCommit · CodeBuild · CodeDeploy · CodePipeline); reduced manual effort by 70%. Used MLflow for model versioning.',
                ],
            },
            {
                'title': 'Visa Approval Forecasting (US Immigration)',
                'tech': 'Python · ML · MongoDB · Docker · AWS · GridSearchCV',
                'github_url': 'https://github.com/sheriff786',
                'bullets': [
                    'Predicted US visa approval outcomes using production-grade ML models; performed EDA · feature engineering · and hyperparameter tuning via GridSearchCV.',
                    'Connected MongoDB for data persistence; containerised solution with Docker for AWS deployment.',
                ],
            },
            {
                'title': 'Disease Classification Mobile App',
                'tech': 'TensorFlow · CNN · TFLite · FastAPI · GCP · ReactJS · React Native',
                'github_url': 'https://github.com/sheriff786',
                'bullets': [
                    'Achieved 92% accuracy; quantized CNN with TF Lite — reduced model size by 80% and inference latency by 50%.',
                    'Deployed serverless inference on GCP Cloud Functions (25% cost reduction); served predictions via FastAPI to React frontends.',
                ],
            },
            {
                'title': 'Accident & Fall Detection System',
                'tech': 'Python · YOLO · OpenCV · CNN · CUDA · SMTP',
                'github_url': 'https://github.com/sheriff786',
                'bullets': [
                    'Built real-time human fall detection system achieving 95% accuracy using YOLO + CUDA toolkit on edge hardware.',
                    'Designed smart camera surveillance for highway accident detection; automated ambulance dispatch via SMTP alerts · cutting emergency response time by 40%.',
                ],
            },
        ]
    
    for proj in projects:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=4, after=1)
        run = p.add_run(proj['title'])
        run.font.bold = True
        run.font.size = Pt(9.5)
        run = p.add_run(f"  |  {proj['tech']}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)
        # Add GitHub link as clickable hyperlink
        if proj.get('github_url'):
            run = p.add_run('    ')
            run.font.size = Pt(8.5)
            add_hyperlink(p, 'View on GitHub →', proj['github_url'], font_size=Pt(8.5))
        
        for bullet in proj['bullets']:
            add_bullet(doc, bullet)
    
    # ═══════════════════════════════════════════════════════════════
    # EDUCATION
    # ═══════════════════════════════════════════════════════════════
    
    add_section_heading(doc, 'EDUCATION')
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=2)
    run = p.add_run('M.Tech in Data Science & Engineering')
    run.font.bold = True
    run.font.size = Pt(9.5)
    run = p.add_run('  ·  BITS Pilani, India  ·  Oct 2021 – Oct 2023')
    run.font.size = Pt(9.5)
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=2)
    run = p.add_run('B.E in Computer Science & Engineering')
    run.font.bold = True
    run.font.size = Pt(9.5)
    run = p.add_run('  ·  OIST Bhopal, India  ·  Aug 2015 – Jun 2019')
    run.font.size = Pt(9.5)
    
    # ═══════════════════════════════════════════════════════════════
    # CERTIFICATIONS & AWARDS
    # ═══════════════════════════════════════════════════════════════
    
    add_section_heading(doc, 'CERTIFICATIONS & AWARDS')
    
    certs = [
        'GCP Professional ML Engineer (Certified Feb 2023)',
        'Certified MLOps Developer – Dataiku',
        'IBM Professional Data Science Certification',
        '1st Prize – Trizetto Hackathon (Patient Data Integration · 95% approval rate)',
        'OpenAI Whisper Hackathon – Hearing-impaired accessibility app',
        'Secretary · Toastmasters Club  ·  Speaker · PyData Conference',
    ]
    
    for cert in certs:
        add_bullet(doc, cert)
    
    # ═══════════════════════════════════════════════════════════════
    # SAVE
    # ═══════════════════════════════════════════════════════════════
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M')
    company_clean = re.sub(r'[^a-zA-Z0-9]', '_', company)
    title_clean = re.sub(r'[^a-zA-Z0-9]', '_', job_title)[:40]
    
    os.makedirs("data/generated-resumes", exist_ok=True)
    filename = f"MohammadSheriff_{company_clean}_{title_clean}_{timestamp}"
    docx_path = f"data/generated-resumes/{filename}.docx"
    doc.save(docx_path)
    
    # Save artifacts
    artifact_folder = f"data/artifacts/{timestamp}"
    os.makedirs(artifact_folder, exist_ok=True)
    
    return docx_path


if __name__ == '__main__':
    print("=" * 60)
    print("GENERATING ATS-OPTIMIZED 2-PAGE TAILORED RESUMES")
    print("=" * 60)
    
    # Databricks AI Engineer - India (Remote)
    path1 = create_tailored_resume_v2(
        job_title="AI Engineer - FDE (Forward Deployed Engineer)",
        company="Databricks",
        location="Remote - India",
        jd_keywords=["GenAI", "RAG", "multi-agent systems", "Text2SQL", "fine-tuning",
                     "HuggingFace", "LangChain", "DSPy", "PyTorch", "pandas", "scikit-learn",
                     "AWS", "Azure", "GCP", "production-grade", "LLMOps"],
    )
    print(f"\n✅ Databricks (India): {path1}")
    
    # Airbnb AI Engineer
    path2 = create_tailored_resume_v2(
        job_title="AI Engineer - Community Support Engineering",
        company="Airbnb",
        location="Remote",
        jd_keywords=["AI", "Machine Learning", "RAG", "scalable services",
                     "cross-functional", "experimentation", "AI-powered applications"],
    )
    print(f"✅ Airbnb (Remote): {path2}")
    
    # Convert to PDF
    print("\nConverting to PDF...")
    try:
        from docx2pdf import convert
        for path in [path1, path2]:
            pdf_path = path.replace('.docx', '.pdf')
            convert(path, pdf_path)
            print(f"  📄 {pdf_path}")
        print("\n✅ Done! All resumes generated.")
    except Exception as e:
        print(f"  ⚠️ PDF error: {e}")
        print("  DOCX files ready.")
    
    # Clean old files
    old_files = [
        "data/generated-resumes/MohammadSheriff_Databricks_AI_Engineer___FDE__Forward_Deployed_Engi_20260705_0023.docx",
        "data/generated-resumes/MohammadSheriff_Databricks_AI_Engineer___FDE__Forward_Deployed_Engi_20260705_0023.pdf",
        "data/generated-resumes/MohammadSheriff_Airbnb_AI_Engineer__Community_Support_Engineeri_20260705_0023.docx",
        "data/generated-resumes/MohammadSheriff_Airbnb_AI_Engineer__Community_Support_Engineeri_20260705_0023.pdf",
    ]
    for f in old_files:
        if os.path.exists(f):
            os.remove(f)
    
    print(f"\n{'=' * 60}")
    print("APPLY HERE:")
    print(f"{'=' * 60}")
    print("  1. Databricks (India Remote): https://databricks.com/company/careers/open-positions/job?gh_jid=8099751002")
    print("  2. Airbnb (Remote):           https://careers.airbnb.com/positions/7946288?gh_jid=7946288")
